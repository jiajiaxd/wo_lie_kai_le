from docx import Document

# 打开 Word 文档
doc = Document('tk.docx')  # 替换为你的文件路径


def has_highlight(paragraph):
    """
    检查段落是否有颜色覆盖标注
    """
    # 将段落的 XML 转化为字符串，检查 highlight 属性
    para_xml = paragraph._p.xml  # 获取段落的底层 XML
    return 'w:highlight' in para_xml  # 判断 XML 中是否包含高亮标注


def del_question_number(text):
    # 先对前面几个字逐个检查，直到找到第一个不是数字的字符
    for j in range(len(text)):
        if not text[j].isdigit():
            break
    # 如果第一个不是数字的字符是.或者空格，那么去除
    if text[j] == '.' or text[j] == ' ':
        text = text[j + 1 :]
    else:
        text = text[j:]
    return text


tiku = {"multiple_choice": [], "judgement": []}

# 逐段读取 Word 文档的内容
i = -1
for paragraph in doc.paragraphs:
    i += 1
    # 如果是空白段落，跳过
    if paragraph.text.strip() == '':
        continue
    # 先判断题型
    # 如果去除空格后，末尾是Y或者N，那么这个paragraph就是判断题
    if paragraph.text.strip()[-1] == 'Y' or paragraph.text.strip()[-1] == 'N':
        judgement = {'question': '', 'answer': False}
        judgement['question'] = del_question_number(paragraph.text)
        # 问题还要删除末尾的Y或者N
        judgement['question'] = judgement['question'][:-1]
        # 如果末尾是Y，那么答案是True 否则是False
        if paragraph.text.strip()[-1] == 'Y':
            judgement['answer'] = True
        tiku['judgement'].append(judgement)
        print(judgement)
        continue

    # 如果下一个不是空白的paragraph的字符串带有A，下下个带有B..，那么这个paragraph就是题干
    step = 0
    next_four_not_blank = []
    while step < 4:
        step += 1
        # 防止报错，如果超出范围，直接跳出
        if i + step >= len(doc.paragraphs):
            step = -1
            break
        if doc.paragraphs[i + step].text.strip() != '':
            next_four_not_blank.append(doc.paragraphs[i + step])
    if step == -1:
        continue
    if len(next_four_not_blank) == 4:
        if (
            'A' in next_four_not_blank[0].text
            and 'B' in next_four_not_blank[1].text
            and 'C' in next_four_not_blank[2].text
            and 'D' in next_four_not_blank[3].text
        ):
            question = {'question': '', 'choices': [], 'answer': ''}
            # 题干需要去除题号，题号一定包含数字，但是不一定包含.，所以直接去除数字，如果去除数字后第一个字符是.，那么去除，如果是空格，那么去除
            question['question'] = del_question_number(paragraph.text)

            for choice in next_four_not_blank:
                # 规范化选项格式
                # 目前文档的选项格式的括号既有中文也有英文还有混合的
                # 为了方便处理，这里统一将括号替换为英文括号
                choice_text = choice.text.replace('（', '(').replace('）', ')')
                # 选项包含在括号中，所以通过括号来切分选项
                choice_text = choice_text.split(')')[1].strip()
                question['choices'].append(choice_text)
                if has_highlight(choice):
                    question['answer'] = choice_text
            print(question)
            tiku['multiple_choice'].append(question)
# 将题库保存为 JSON 文件
import json

with open('tiku.json', 'w', encoding='utf-8') as f:
    json.dump(tiku, f, ensure_ascii=False, indent=2)
print(f'题库已经提取完成。共提取了{len(tiku["multiple_choice"])}道选择题，{len(tiku["judgement"])}道判断题。\nHave Fun!')

with open('tiku.md', 'w', encoding='utf-8') as f:
    f.write('# 题库\n\n')
    f.write('## 选择题\n\n')
    for i, q in enumerate(tiku['multiple_choice'], 1):
        f.write(f"### 问题 {i}\n\n")
        f.write(f"**题目**: {q['question']}\n\n")
        f.write("**选项**:  \n")
        for j, opt in enumerate(q['choices']):
            f.write(f"{chr(65+j)}. {opt}  \n")
        f.write('\n')
        f.write('<details>\n<summary>查看答案</summary>\n\n')
        f.write(f"{chr(65 + q['choices'].index(q['answer']))}. {q['answer']}\n\n")
        f.write('</details>\n\n')
    f.write('## 判断题\n\n')
    for i, q in enumerate(tiku['judgement'], 1):
        f.write(f"### 问题 {i}\n\n")
        f.write(f"**题目**: {q['question']}\n\n")
        f.write('<details>\n<summary>查看答案</summary>\n\n')
        f.write(f"{q['answer']}\n\n")
        f.write('</details>\n\n')